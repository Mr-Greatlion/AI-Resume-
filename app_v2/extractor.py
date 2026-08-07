"""
Extraction layer for v2.

Priority order:
  1. Import the existing v1 extractors from app/ (best accuracy — same code v1 uses)
  2. If v1 not importable (different server path), fall back to built-in regex

This means v2 accuracy is identical to v1 when deployed on the same VPS.
"""
from __future__ import annotations
import os, re, sys, time, tempfile
from datetime import date


# ── helpers ───────────────────────────────────────────────────────────────────

EMAIL_RE  = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
PHONE_RE  = re.compile(r"(?:\+?91[\-\s]?)?[6-9]\d{9}\b")
PAN_RE    = re.compile(r"\b[A-Z]{5}[0-9]{4}[A-Z]\b")
AAD_RE    = re.compile(r"\b\d{4}[\s\-]?\d{4}[\s\-]?\d{4}\b")
PIN_RE    = re.compile(r"\b[1-9]\d{5}\b")
EXP_RE    = re.compile(r"(\d{1,2})\s*\+?\s*(?:years?|yrs?)", re.I)
NAME_BAD  = re.compile(r"\d|@|resume|curriculum|profile|summary|objective", re.I)


def _exp_bucket(years: int) -> str:
    if years < 1:  return "Fresher"
    if years <= 1: return "0-1 years"
    if years <= 3: return "1-3 years"
    if years <= 5: return "3-5 years"
    if years <= 10:return "5-10 years"
    return "10+ years"


def _read_text(path: str, filename: str) -> tuple[str, str]:
    """Return (text, method). method = pdf | pdf-ocr | docx | image | txt"""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        try:
            from pdfminer.high_level import extract_text
            t = extract_text(path) or ""
            if t.strip():
                return t[:20000], "pdf"
        except Exception:
            pass
        # scanned — OCR
        try:
            import easyocr, pdf2image
            reader = easyocr.Reader(["en"], gpu=False)
            imgs = pdf2image.convert_from_path(path, dpi=200)
            parts = []
            for img in imgs:
                tmp = tempfile.mktemp(suffix=".jpg")
                img.save(tmp)
                parts.extend(reader.readtext(tmp, detail=0))
                os.remove(tmp)
            return " ".join(parts)[:20000], "pdf-ocr"
        except Exception as e:
            return "", "pdf-ocr"

    if ext == ".docx":
        try:
            import docx as _docx
            doc = _docx.Document(path)
            lines = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    lines.append(" ".join(c.text for c in row.cells))
            return "\n".join(lines)[:20000], "docx"
        except Exception:
            return "", "docx"

    if ext in (".jpg", ".jpeg", ".png"):
        try:
            import easyocr
            reader = easyocr.Reader(["en"], gpu=False)
            return " ".join(reader.readtext(path, detail=0))[:20000], "image"
        except Exception:
            return "", "image"

    if ext == ".txt":
        for enc in ("utf-8", "utf-16", "latin-1"):
            try:
                return open(path, encoding=enc).read()[:20000], "txt"
            except Exception:
                pass

    return "", "unknown"


def _regex_extract(text: str) -> dict:
    """Pure-regex fallback — works with zero ML dependencies."""
    low = text.lower()

    emails = sorted(set(EMAIL_RE.findall(text)))
    phones = sorted(set(
        "+91 " + re.sub(r"\D","",m)[-10:]
        for m in PHONE_RE.findall(text)
        if len(re.sub(r"\D","",m)) >= 10
    ))

    pan = next(iter(PAN_RE.findall(text)), "")

    aadhaar = ""
    for m in AAD_RE.finditer(text):
        d = re.sub(r"\D","",m.group(0))
        if len(d) == 12 and d[0] not in "01":
            aadhaar = f"{d[:4]} {d[4:8]} {d[8:]}"; break

    pincode = next((p for p in PIN_RE.findall(text)
                    if p not in "".join(phones)), "")

    exp_yrs = [int(m.group(1)) for m in EXP_RE.finditer(text) if int(m.group(1)) <= 50]
    experience = _exp_bucket(max(exp_yrs)) if exp_yrs else (
        "Fresher" if "fresher" in low else "")

    name = ""
    for line in text.splitlines():
        l = line.strip()
        if 3 <= len(l) <= 40 and not NAME_BAD.search(l):
            if re.fullmatch(r"[A-Za-z][A-Za-z .\-']+", l) and 1 <= len(l.split()) <= 4:
                name = l.title(); break

    edu = ""
    for kw, val in [
        ("ph.d","Doctorate"),("mba","PG / MBA"),("m.tech","PG / M.Tech"),
        ("master","PG (Master)"),("b.e","UG (B.E)"),("b.tech","UG (B.Tech)"),
        ("bachelor","UG (Bachelor)"),("diploma","Diploma"),
        ("12th","12th / HSC"),("hsc","12th / HSC"),
        ("10th","10th / SSLC"),("sslc","10th / SSLC")
    ]:
        if re.search(r"(?<![a-z])" + re.escape(kw) + r"(?![a-z])", low):
            edu = val; break

    city = state = ""
    for c_name in ["Chennai","Coimbatore","Madurai","Bengaluru","Bangalore",
                   "Hyderabad","Mumbai","Delhi","Pune","Trichy"]:
        if c_name.lower() in low: city = c_name; break
    for s_name in ["Tamil Nadu","Tamilnadu","Karnataka","Maharashtra",
                   "Andhra Pradesh","Telangana","Kerala"]:
        if s_name.lower() in low: state = s_name; break

    return {
        "name": name, "job_title": "", "education": edu,
        "experience": experience, "emails": emails, "phones": phones,
        "address": "", "city": city, "state": state,
        "country": "India" if "india" in low else "",
        "pincode": pincode, "pan": pan, "aadhaar": aadhaar,
        "current_location": f"{city}, {state}".strip(", "),
        "is_employee": "Candidate",
        "applied_date": date.today().isoformat(),
    }


def _v1_extract(text: str) -> dict | None:
    """Try to use v1 extractors — same accuracy as the live v1 API."""
    from app_v2.config import V1_APP_PATH
    if V1_APP_PATH not in sys.path:
        sys.path.insert(0, V1_APP_PATH)
    try:
        from app.extractor import (
            extract_name, extract_email, extract_phone,
            extract_job_title, detect_pan, detect_aadhaar,
        )
        from app.education_extractor import extract_education
        from app.experience_calc import calculate_experience
        from app.address_extractor import extract_address, extract_current_location

        emails  = sorted(set(extract_email(text[:3000])))
        phones  = sorted(set(extract_phone(text[:3000])))
        addr    = extract_address(text) or {}
        loc     = extract_current_location(text) or {}
        loc_str = ", ".join(x for x in [loc.get("city",""),loc.get("state","")] if x)

        return {
            "name":             extract_name(text) or "",
            "job_title":        extract_job_title(text) or "",
            "education":        extract_education(text) or "",
            "experience":       calculate_experience(text),
            "emails":           emails,
            "phones":           phones,
            "address":          addr.get("address",""),
            "city":             addr.get("city",""),
            "state":            addr.get("state",""),
            "country":          addr.get("country","India"),
            "pincode":          addr.get("pincode",""),
            "pan":              detect_pan(text) or "",
            "aadhaar":          detect_aadhaar(text) or "",
            "current_location": loc_str,
            "is_employee":      "Candidate",
            "applied_date":     date.today().isoformat(),
        }
    except ImportError:
        return None


def run(path: str, filename: str) -> tuple[dict, str, str, int, str | None]:
    """
    Main entry point.
    Returns (fields, proc_status, extract_method, duration_ms, error)
    proc_status: success | partial | failed
    """
    t0 = time.perf_counter()
    error: str | None = None

    text, method = _read_text(path, filename)
    if not text.strip():
        ms = int((time.perf_counter()-t0)*1000)
        return {}, "failed", method, ms, "Could not extract text from file"

    # try v1 extractors first, fall back to regex
    try:
        fields = _v1_extract(text) or _regex_extract(text)
    except Exception as e:
        error = f"Extractor error: {e}"
        fields = _regex_extract(text)

    ms = int((time.perf_counter()-t0)*1000)

    # status
    has_name  = bool(fields.get("name"))
    has_email = bool(fields.get("emails"))
    has_phone = bool(fields.get("phones"))
    if not any([has_name, has_email, has_phone]):
        status = "failed"
    elif not all([has_name, has_email, has_phone]):
        status = "partial"
    else:
        status = "success"

    return fields, status, method, ms, error
