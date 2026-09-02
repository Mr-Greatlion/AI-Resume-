"""
File → plain text for v2.

Supported:  .pdf (text layer, OCR fallback)  .docx  .doc (antiword → LibreOffice)
            .rtf  .txt  .jpg/.jpeg/.png/.webp/.bmp/.tif/.tiff (OCR)

OCR engine order: EasyOCR (if installed) → pytesseract (if installed + binary).
Every reader is wrapped so a missing optional library can never crash the API —
the caller just gets ("", method) and reports "Could not extract text".
"""
from __future__ import annotations

import logging
import os
import re
import shutil
import subprocess
import tempfile

from app_v2.config import MAX_TEXT_LENGTH, OCR_ENABLED, OCR_GPU

log = logging.getLogger("resume_v2.textract")

_easyocr_reader = None


def _clean(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_TEXT_LENGTH]


# ── OCR ──────────────────────────────────────────────────────────────────────

def _ocr_image(path: str) -> str:
    """OCR one image file. Returns "" when no OCR engine is available."""
    if not OCR_ENABLED:
        return ""
    global _easyocr_reader
    try:
        import easyocr  # type: ignore
        if _easyocr_reader is None:
            _easyocr_reader = easyocr.Reader(["en"], gpu=OCR_GPU, verbose=False)
        parts = _easyocr_reader.readtext(path, detail=0, paragraph=True)
        text = "\n".join(p for p in parts if p)
        if text.strip():
            return text
    except Exception as exc:  # library missing or model download failed
        log.debug("easyocr unavailable/failed: %s", exc)
    try:
        import pytesseract  # type: ignore
        from PIL import Image
        return pytesseract.image_to_string(Image.open(path))
    except Exception as exc:
        log.debug("pytesseract unavailable/failed: %s", exc)
    return ""


def _pdf_to_images(path: str, dpi: int = 220) -> list[str]:
    """Rasterise a PDF into temp JPGs. Uses pdf2image if present, else pdftoppm CLI."""
    out: list[str] = []
    tmpdir = tempfile.mkdtemp(prefix="resume_v2_ocr_")
    try:
        from pdf2image import convert_from_path  # type: ignore
        for i, img in enumerate(convert_from_path(path, dpi=dpi)):
            p = os.path.join(tmpdir, f"page_{i}.jpg")
            img.save(p, "JPEG")
            out.append(p)
        return out
    except Exception as exc:
        log.debug("pdf2image failed: %s", exc)
    if shutil.which("pdftoppm"):
        try:
            subprocess.run(["pdftoppm", "-r", str(dpi), "-jpeg", path, os.path.join(tmpdir, "page")],
                           check=True, capture_output=True, timeout=120)
            out = sorted(os.path.join(tmpdir, f) for f in os.listdir(tmpdir) if f.endswith(".jpg"))
        except Exception as exc:
            log.debug("pdftoppm failed: %s", exc)
    return out


# ── Readers ──────────────────────────────────────────────────────────────────

def _read_pdf(path: str) -> tuple[str, str]:
    text = ""
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(path) or ""
    except Exception as exc:
        log.warning("pdfminer failed: %s", exc)
    if len(text.strip()) >= 80:
        return text, "pdf"
    # scanned PDF → OCR
    pages = _pdf_to_images(path)
    ocr_parts = []
    for p in pages:
        try:
            ocr_parts.append(_ocr_image(p))
        finally:
            try:
                os.remove(p)
            except OSError:
                pass
    if pages:
        shutil.rmtree(os.path.dirname(pages[0]), ignore_errors=True)
    ocr_text = "\n".join(t for t in ocr_parts if t)
    if len(ocr_text.strip()) > len(text.strip()):
        return ocr_text, "pdf-ocr"
    return text, "pdf"


def _read_docx(path: str) -> tuple[str, str]:
    try:
        import docx  # type: ignore
        d = docx.Document(path)
        lines: list[str] = []
        for sec in d.sections:
            for p in sec.header.paragraphs:
                if p.text.strip():
                    lines.append(p.text)
        lines.extend(p.text for p in d.paragraphs)
        for tbl in d.tables:
            for row in tbl.rows:
                cells = []
                for c in row.cells:
                    t = c.text.strip()
                    if t and t not in cells:
                        cells.append(t)
                if cells:
                    lines.append(" | ".join(cells))
        text = "\n".join(lines)
        if text.strip():
            return text, "docx"
    except Exception as exc:
        log.warning("python-docx failed: %s", exc)
    try:
        import docx2txt  # type: ignore
        return docx2txt.process(path) or "", "docx"
    except Exception:
        pass
    return _read_via_libreoffice(path, "docx")


def _read_via_libreoffice(path: str, method: str) -> tuple[str, str]:
    exe = shutil.which("libreoffice") or shutil.which("soffice")
    if not exe:
        return "", method
    outdir = tempfile.mkdtemp(prefix="resume_v2_lo_")
    try:
        subprocess.run([exe, "--headless", "--convert-to", "txt:Text", "--outdir", outdir, path],
                       check=True, capture_output=True, timeout=120)
        for f in os.listdir(outdir):
            if f.endswith(".txt"):
                with open(os.path.join(outdir, f), encoding="utf-8", errors="ignore") as fh:
                    return fh.read(), method
    except Exception as exc:
        log.warning("libreoffice convert failed: %s", exc)
    finally:
        shutil.rmtree(outdir, ignore_errors=True)
    return "", method


def _read_doc(path: str) -> tuple[str, str]:
    if shutil.which("antiword"):
        try:
            r = subprocess.run(["antiword", path], capture_output=True, timeout=60)
            if r.returncode == 0 and r.stdout.strip():
                return r.stdout.decode("utf-8", errors="ignore"), "doc"
        except Exception as exc:
            log.debug("antiword failed: %s", exc)
    return _read_via_libreoffice(path, "doc")


def _read_rtf(path: str) -> tuple[str, str]:
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
        with open(path, encoding="utf-8", errors="ignore") as fh:
            return rtf_to_text(fh.read()), "rtf"
    except Exception:
        return _read_via_libreoffice(path, "rtf")


def _read_txt(path: str) -> tuple[str, str]:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            with open(path, encoding=enc) as fh:
                return fh.read(), "txt"
        except Exception:
            continue
    return "", "txt"


def _read_image(path: str) -> tuple[str, str]:
    src = path
    tmp = None
    try:
        from PIL import Image
        img = Image.open(path)
        if img.mode not in ("RGB", "L") or not path.lower().endswith((".jpg", ".jpeg", ".png")):
            tmp = tempfile.mktemp(suffix=".png", prefix="resume_v2_img_")
            img.convert("RGB").save(tmp, "PNG")
            src = tmp
    except Exception as exc:
        log.debug("PIL normalise failed: %s", exc)
    try:
        return _ocr_image(src), "image"
    finally:
        if tmp and os.path.exists(tmp):
            os.remove(tmp)


# ── Public entry ─────────────────────────────────────────────────────────────

def file_to_text(path: str, filename: str | None = None) -> tuple[str, str]:
    """Return (clean_text, method). Never raises."""
    ext = os.path.splitext(filename or path)[1].lower()
    try:
        if ext == ".pdf":
            text, method = _read_pdf(path)
        elif ext == ".docx":
            text, method = _read_docx(path)
        elif ext == ".doc":
            text, method = _read_doc(path)
        elif ext == ".rtf":
            text, method = _read_rtf(path)
        elif ext == ".txt":
            text, method = _read_txt(path)
        elif ext in (".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff", ".gif"):
            text, method = _read_image(path)
        else:
            return "", "unsupported"
    except Exception as exc:
        log.exception("file_to_text crashed: %s", exc)
        return "", ext.lstrip(".") or "unknown"
    return _clean(text), method
