"""
End-to-end tests for Resume Extractor v2.
Run:  pytest tests_v2 -q
Uses a temporary SQLite DB + upload folders — never touches production data.
"""
from __future__ import annotations

import io
import json
import os
import sys
import tempfile

import pytest

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)

TMP = tempfile.mkdtemp(prefix="resume_v2_test_")
os.environ["V2_BASE_DIR"] = TMP
os.environ["V2_UPLOAD_DIR"] = os.path.join(TMP, "uploads")
os.environ["V2_CERT_DIR"] = os.path.join(TMP, "certs")
os.environ["DATABASE_URL"] = "sqlite:///" + os.path.join(TMP, "test.db")
os.environ["RESUME_V2_KEY"] = "test_key_123"
os.environ["MAX_RECORDS"] = "0"
os.environ["JOB_TITLE_TTL"] = "1"

from fastapi.testclient import TestClient  # noqa: E402

from app_v2.config import API_PREFIX  # noqa: E402
from app_v2.main import app  # noqa: E402

SAMPLE_PDF = os.path.join(ROOT, "data", "sample_resume.pdf")
KEY = {"x-api-key": "test_key_123"}

NIGERIAN_TXT = """CI CALISTUS IWUJI
07034577995   tochicalistus4real@gmail.com
09152390557

PROFESSIONAL Result oriented, technically sound and experience in Electrical and Instrumental Technician for about 10 years demonstrated history of working in oil and gas sectors. Passionate technician with HND in Electrical Electronics Engineering power system and Instrumentation Engineering from Akanu Ibiam Federal Polytechnic, Unwana, Ebonyi State.

WORK EXPERIENCE
Electrical Technician - Shell Petroleum Development Company, Port Harcourt, Rivers State
2016 - Present

EDUCATION
Higher National Diploma (HND) Electrical/Electronics Engineering, Akanu Ibiam Federal Polytechnic, Unwana 2014

CERTIFICATIONS
HSE Level 3, Confined Space Entry, First Aid

ADDRESS: No 12 Aba Road, Umuahia, Abia State
"""

INDIAN_TXT = """PRADEEP KUMAR R
Boiler Operator
Mobile: +91 98400 12345 | Email: pradeep.kumar@example.com
Address: No. 14, Gandhi Nagar, Anna Nagar West, Chennai - 600040, Tamil Nadu, India
PAN: ABCPK1234F   Aadhaar: 4321 8765 1234

CAREER OBJECTIVE
Boiler Operator with 4 years of experience in AFBC boiler operation at a 30 MW power plant.

EXPERIENCE
Boiler DCS Operator - ABC Power Ltd, Coimbatore  (Jan 2021 - Present)
Boiler Field Operator - XYZ Sugars, Erode (Jun 2019 - Dec 2020)

EDUCATION
Diploma in Mechanical Engineering, Govt Polytechnic Chennai, 2019
First Class Boiler Attendant certificate
"""


@pytest.fixture(scope="module")
def client():
    with TestClient(app) as c:
        yield c


def _upload(client, name: str, data: bytes, mime="application/octet-stream"):
    return client.post(f"{API_PREFIX}/extract", files={"file": (name, io.BytesIO(data), mime)})


# ── 1. health + options ──────────────────────────────────────────────────────

def test_health(client):
    r = client.get(f"{API_PREFIX}/health")
    assert r.status_code == 200
    d = r.json()
    assert d["status"] == "running" and d["dbOk"] is True and d["database"] == "sql"


def test_options(client):
    d = client.get(f"{API_PREFIX}/options").json()
    assert len(d["jobTitles"]) > 20
    assert "5-10 years" in d["experience"]
    assert "HND" in d["education"]
    assert "Nigeria" in d["countries"] and "Ebonyi" in d["states"]["Nigeria"]
    assert "+234" in d["countryCodes"]


# ── 2. extraction ────────────────────────────────────────────────────────────

def test_extract_sample_pdf(client):
    with open(SAMPLE_PDF, "rb") as fh:
        r = _upload(client, "sample_resume.pdf", fh.read(), "application/pdf")
    assert r.status_code == 200, r.text
    d = r.json()
    f = d["fields"]
    assert d["status"] == "success" and len(d["uid"]) == 32
    assert f["fullName"]["value"] == "Karanraj Duraikannu" and f["fullName"]["status"] == "extracted"
    assert f["emails"]["value"][0]["emailAddress"] == "d.karanraj1997@gmail.com"
    assert f["mobileNumbers"]["value"][0] == {"countryCode": "+91", "mobileNumber": "9600901728", "isPrimary": True}
    assert f["jobTitle"]["value"] == "Software Engineer"
    assert f["yearsOfExperience"]["value"] == "5-10 years"
    assert f["educationQualification"]["value"] == "B.E"
    assert f["currentWorkLocation"]["value"]["city"] == "Bangalore"
    assert f["currentWorkLocation"]["value"]["state"] == "Karnataka"
    assert d["meta"]["detectedCountry"] == "India"
    assert d["fileUrl"].endswith(d["uid"])


def test_extract_nigerian_txt(client):
    r = _upload(client, "calistus.txt", NIGERIAN_TXT.encode(), "text/plain")
    assert r.status_code == 200, r.text
    f = r.json()["fields"]
    assert f["fullName"]["value"] == "Calistus Iwuji"          # "CI" logo junk removed
    assert f["surname"]["value"] == "Iwuji"
    phones = f["mobileNumbers"]["value"]
    assert [p["mobileNumber"] for p in phones] == ["7034577995", "9152390557"]
    assert all(p["countryCode"] == "+234" for p in phones)      # NOT +91
    assert f["jobTitle"]["value"] in ("Electrical Technician", "Instrumentation Technician")
    assert f["jobTitle"]["matchedOption"] is True
    assert f["yearsOfExperience"]["value"] == "5-10 years"
    assert f["educationQualification"]["value"] == "HND"
    assert {c["name"] for c in f["certificates"]["value"]} >= {"HSE Level 3", "Confined Space Entry", "First Aid / CPR"}
    cur = f["currentWorkLocation"]["value"]
    assert cur["country"] == "Nigeria" and cur["state"] == "Rivers" and cur["city"] == "Port Harcourt"
    addr = f["permanentAddress"]["value"]
    assert addr["address"].startswith("No 12 Aba Road") and addr["city"] == "Umuahia" and addr["state"] == "Abia"
    assert addr["country"] == "Nigeria" and addr["pinCode"] == ""
    assert r.json()["meta"]["defaultDialCode"] == "+234"


def test_extract_indian_txt(client):
    r = _upload(client, "pradeep.txt", INDIAN_TXT.encode(), "text/plain")
    assert r.status_code == 200, r.text
    f = r.json()["fields"]
    assert f["fullName"]["value"] == "Pradeep Kumar R"
    assert f["mobileNumbers"]["value"][0]["mobileNumber"] == "9840012345"
    assert f["mobileNumbers"]["value"][0]["countryCode"] == "+91"
    assert f["jobTitle"]["value"] in ("Boiler Operator", "Boiler DCS Operator")
    assert f["yearsOfExperience"]["value"] == "3-5 years"
    assert f["educationQualification"]["value"] == "Diploma (Mechanical)"
    assert f["pan"]["value"] == "ABCPK1234F"
    assert f["aadhar"]["value"] == "432187651234"
    addr = f["permanentAddress"]["value"]
    assert addr["city"] == "Chennai" and addr["state"] == "Tamil Nadu" and addr["pinCode"] == "600040"
    assert addr["address"].startswith("No. 14, Gandhi Nagar")
    assert "First Class Boiler Attendant" in {c["name"] for c in f["certificates"]["value"]}


def test_extract_docx(client):
    import docx
    d = docx.Document()
    d.add_paragraph("ARUN PRASAD")
    d.add_paragraph("Instrumentation Engineer")
    d.add_paragraph("arun.prasad@gmail.com | 9876543210")
    d.add_paragraph("Instrumentation Engineer with 6 years of experience in DCS and PLC systems, Hyderabad, Telangana.")
    d.add_paragraph("Education: B.Tech Electronics and Instrumentation, JNTU 2018")
    buf = io.BytesIO()
    d.save(buf)
    r = _upload(client, "arun.docx", buf.getvalue())
    assert r.status_code == 200, r.text
    f = r.json()["fields"]
    assert r.json()["extractMethod"] == "docx"
    assert f["fullName"]["value"] == "Arun Prasad"
    assert f["jobTitle"]["value"] == "Instrumentation Engineer"
    assert f["educationQualification"]["value"] == "B.Tech"
    assert f["yearsOfExperience"]["value"] == "5-10 years"
    assert f["currentWorkLocation"]["value"]["city"] == "Hyderabad"


def test_extract_rejects_bad_input(client):
    assert _upload(client, "x.exe", b"abc").status_code == 400
    assert _upload(client, "empty.pdf", b"").status_code == 400
    assert _upload(client, "big.pdf", b"0" * (11 * 1024 * 1024)).status_code == 413
    r = _upload(client, "junk.pdf", b"not a real pdf")
    assert r.status_code == 200 and r.json()["status"] == "failed" and r.json()["error"]


# ── 3. file preview + certificates ───────────────────────────────────────────

def test_file_preview_and_certificate(client):
    with open(SAMPLE_PDF, "rb") as fh:
        pdf = fh.read()
    uid = _upload(client, "sample_resume.pdf", pdf, "application/pdf").json()["uid"]
    r = client.get(f"{API_PREFIX}/files/{uid}")
    assert r.status_code == 200 and r.headers["content-type"].startswith("application/pdf")
    assert r.content == pdf and "inline" in r.headers["content-disposition"]
    assert client.get(f"{API_PREFIX}/files/doesnotexist").status_code == 404

    r = client.post(f"{API_PREFIX}/certificates/{uid}", data={"name": "NEBOSH IGC"},
                    files={"file": ("nebosh.pdf", io.BytesIO(pdf), "application/pdf")})
    assert r.status_code == 200, r.text
    cert = r.json()
    assert cert["name"] == "NEBOSH IGC" and cert["fileUrl"].endswith(cert["fileId"])
    r2 = client.get(cert["fileUrl"])
    assert r2.status_code == 200 and r2.content == pdf
    assert client.post(f"{API_PREFIX}/certificates/badUID", files={"file": ("a.pdf", io.BytesIO(pdf))}).status_code == 404
    assert client.post(f"{API_PREFIX}/certificates/{uid}", files={"file": ("a.exe", io.BytesIO(b"x"))}).status_code == 400


# ── 4. review → stored in DB ─────────────────────────────────────────────────

def test_review_roundtrip(client):
    uid = _upload(client, "calistus.txt", NIGERIAN_TXT.encode(), "text/plain").json()["uid"]
    reviewed = {
        "fullName": "Calistus Iwuji", "surname": "Iwuji",
        "emails": [{"emailAddress": "tochicalistus4real@gmail.com", "isPrimary": True}],
        "mobileNumbers": [{"countryCode": "+234", "mobileNumber": "7034577995", "isPrimary": True},
                          {"countryCode": "+234", "mobileNumber": "9152390557", "isPrimary": False}],
        "jobTitle": "Instrumentation Technician",              # reviewer changes this
        "yearsOfExperience": "10+ years",                       # and this
        "educationQualification": "HND",
        "certificates": [{"name": "HSE Level 3", "fileId": None}],
        "currentWorkLocation": {"country": "Nigeria", "state": "Rivers", "city": "Port Harcourt"},
        "permanentAddress": {"address": "No 12 Aba Road, Umuahia", "city": "Umuahia", "state": "Abia", "country": "Nigeria", "pinCode": ""},
        "pan": "", "aadhar": "",
    }
    r = client.post(f"{API_PREFIX}/review", json={"uid": uid, "fields": reviewed, "reviewer": "sharfudeen"})
    assert r.status_code == 200, r.text
    d = r.json()
    assert d["ok"] and d["reviewId"] >= 1 and d["storedIn"] == "sql"
    assert "jobTitle" in d["changedFields"] or "yearsOfExperience" in d["changedFields"]
    assert "fullName" not in d["changedFields"] and "emails" not in d["changedFields"]
    assert 0 < d["extractionAccuracyPct"] < 100
    assert d["candidate"]["isEmployee"] == "candidate" and d["candidate"]["resume"].endswith(uid)

    # stored? admin fetch
    rec = client.get(f"{API_PREFIX}/records/{uid}", headers=KEY).json()
    assert rec["reviewStatus"] == "reviewed"
    assert len(rec["reviews"]) == 1
    rv = rec["reviews"][0]
    assert rv["reviewer"] == "sharfudeen"
    assert rv["reviewed"]["jobTitle"] == "Instrumentation Technician"
    assert rv["corrections"]["fullName"]["action"] == "confirmed"
    assert rv["corrections"]["yearsOfExperience"]["action"] == "corrected"
    assert rv["corrections"]["yearsOfExperience"]["extracted"] == "5-10 years"
    assert rv["corrections"]["pan"]["action"] == "confirmed"      # empty → empty

    # second review of the same uid is also stored (history kept)
    reviewed["surname"] = "Iwuji-Okafor"
    r = client.post(f"{API_PREFIX}/review", json={"uid": uid, "fields": reviewed})
    assert r.status_code == 200
    rec = client.get(f"{API_PREFIX}/records/{uid}", headers=KEY).json()
    assert len(rec["reviews"]) == 2 and rec["reviews"][0]["reviewed"]["surname"] == "Iwuji-Okafor"


def test_review_validation(client):
    uid = _upload(client, "x.txt", INDIAN_TXT.encode(), "text/plain").json()["uid"]
    base = {"fullName": "A B", "emails": [{"emailAddress": "a@b.com"}], "jobTitle": "Fitter"}
    assert client.post(f"{API_PREFIX}/review", json={"uid": "nope12345", "fields": base}).status_code == 404
    bad = dict(base, emails=[{"emailAddress": "not-an-email"}])
    assert client.post(f"{API_PREFIX}/review", json={"uid": uid, "fields": bad}).status_code == 422
    bad = dict(base, pan="WRONG")
    assert client.post(f"{API_PREFIX}/review", json={"uid": uid, "fields": bad}).status_code == 422
    bad = dict(base, fullName="")
    assert client.post(f"{API_PREFIX}/review", json={"uid": uid, "fields": bad}).status_code == 422
    bad = dict(base, mobileNumbers=[{"countryCode": "91", "mobileNumber": "98400 12345"}])
    r = client.post(f"{API_PREFIX}/review", json={"uid": uid, "fields": bad})
    assert r.status_code == 200 and r.json()["candidate"]["mobileNumbers"][0] == {"countryCode": "+91", "mobileNumber": "9840012345", "isPrimary": True}
    assert r.json()["candidate"]["permanentAddress"]["country"] == "India"   # default filled


# ── 5. admin ─────────────────────────────────────────────────────────────────

def test_admin_auth_and_lists(client):
    assert client.get(f"{API_PREFIX}/records").status_code == 401
    assert client.get(f"{API_PREFIX}/records", headers={"x-api-key": "wrong"}).status_code == 401
    r = client.get(f"{API_PREFIX}/records", headers=KEY)
    assert r.status_code == 200 and r.json()["total"] >= 5
    assert "rawText" not in r.json()["records"][0]
    assert client.get(f"{API_PREFIX}/records?x_api_key=test_key_123").status_code == 200
    assert client.get(f"{API_PREFIX}/records/nope", headers=KEY).status_code == 404
    s = client.get(f"{API_PREFIX}/stats", headers=KEY).json()
    assert s["extractions"] >= 5 and s["reviews"] >= 3 and s["avgAccuracyPct"] is not None


def test_training_export(client):
    r = client.get(f"{API_PREFIX}/training/export", headers=KEY)
    assert r.status_code == 200
    lines = [json.loads(l) for l in r.text.strip().splitlines()]
    assert len(lines) >= 3
    row = lines[0]
    assert {"uid", "raw_text", "extracted", "reviewed", "corrections"} <= set(row)
    assert row["raw_text"]


def test_ui_and_docs(client):
    r = client.get(f"{API_PREFIX}/")
    assert r.status_code == 200 and "Resume Data Extractor Results" in r.text and "__API_PREFIX__" not in r.text
    assert client.get(f"{API_PREFIX}/docs").status_code == 200
    assert client.get(f"{API_PREFIX}/openapi.json").status_code == 200


# ── 6. retention (MAX_RECORDS) ───────────────────────────────────────────────

def test_retention_purge(monkeypatch):
    from app_v2 import store as st
    monkeypatch.setattr(st, "MAX_RECORDS", 2)
    s = st.SQLStore("sqlite:///" + os.path.join(TMP, "purge.db"))
    paths = []
    for i in range(4):
        p = os.path.join(TMP, f"p{i}.txt")
        open(p, "w").write("x")
        paths.append(p)
        s.save_extraction(dict(uid=f"u{i}", filename=f"f{i}", stored_path=p, extracted={}, meta={}))
    rows, total = s.list_extractions()
    assert total == 2 and [r["uid"] for r in rows] == ["u3", "u2"]
    assert not os.path.exists(paths[0]) and os.path.exists(paths[3])
